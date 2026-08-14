#!/usr/bin/env python3
"""결과보고서 docx 재작성 — 웹사이트 구축·콘텐츠·성과 섹션 채움."""
import json, urllib.request, os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOT = os.path.join(ROOT, "screenshots", "report-2026-08")
DOC = os.path.join(ROOT, "docs", "결과보고서_마케팅PR대행용역_동서남북.docx")

# ── 블로그 데이터 (Supabase REST) ──────────────────────────
env = {}
for line in open(os.path.join(ROOT, ".env.local")):
    if "=" in line and not line.startswith("#"):
        k, v = line.strip().split("=", 1)
        env[k] = v
req = urllib.request.Request(
    env["NEXT_PUBLIC_SUPABASE_URL"] + "/rest/v1/posts"
    "?select=title,category,created_at,views,published&order=created_at.asc",
    headers={"apikey": env["NEXT_PUBLIC_SUPABASE_ANON_KEY"]})
posts = json.load(urllib.request.urlopen(req))
pub = [p for p in posts if p["published"]]
drafts = [p for p in posts if not p["published"]]
total_views = sum(p["views"] for p in pub)
cats = {}
for p in pub:
    cats[p["category"]] = cats.get(p["category"], 0) + 1

# ── 문서 열기, '2. 주요 과업 수행 결과' 이후 삭제 ─────────────
d = docx.Document(DOC)
anchor = next(p for p in d.paragraphs if p.text.strip() == "2. 주요 과업 수행 결과")._element
body = d.element.body
seen = False
for el in list(body):
    if el is anchor:
        seen = True
        continue
    if seen and not el.tag.endswith("sectPr"):
        body.remove(el)

# ── 헬퍼 ──────────────────────────────────────────────────
def H2(text):
    return d.add_paragraph(text, style="Heading 2")

def H3(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    return p

def P(text, bold_prefix=None):
    p = d.add_paragraph()
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)
    return p

def caption(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x6E, 0x7E)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    return p

def picture(fname, cap, width=6.2):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(SHOT, fname), width=Inches(width))
    caption(cap)

def set_borders(table):
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "B9C2D0")
        borders.append(el)
    pr.append(borders)

def table(header, rows, widths=None, size=9):
    t = d.add_table(rows=1, cols=len(header))
    set_borders(t)
    for j, htxt in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = Pt(size)
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), "EEF2FA")
        cell._tc.get_or_add_tcPr().append(sh)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(size)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ══════════════════════════════════════════════════════════
# ① 공식 웹사이트 기획 및 구축
# ══════════════════════════════════════════════════════════
H2("① 공식 웹사이트 기획 및 구축")
P("AI 면접 서비스 홍보를 위한 공식 웹사이트(www.supercoder.co)를 반응형(PC/모바일)으로 기획·디자인하여 구축 완료함. "
  "메인 랜딩과 도입 문의·블로그·법적 고지 등 서브페이지, 그리고 콘텐츠·문의 접수·약관·SEO를 통합 관리하는 "
  "관리자(CMS) 콘솔을 구현함. 검색엔진 등록(구글 서치콘솔·네이버 서치어드바이저·Bing 웹마스터)과 페이지별 "
  "메타데이터·sitemap 등 SEO 기본 최적화, 쿠키 동의 배너 및 개인정보 처리방침 등 컴플라이언스 요소를 반영함.",
  bold_prefix="수행 요약 : ")

H3("1) 일반 사용자 페이지 구성")
table(
    ["페이지", "경로", "주요 내용"],
    [
        ["메인 랜딩", "/", "서비스 선언(히어로) → 고객사 → 핵심 가치 → AI·사람 역할 분담(퍼널) → AI 리포트 소개 → 검증된 효과 → 활용 사례 → FAQ → 상담 신청 CTA로 이어지는 설득형 원페이지 구성"],
        ["도입 문의", "/apply", "도입 상담 신청 폼. 접수 시 DB 저장, 담당자 메일 실시간 알림, GA4 전환 이벤트(apply_lead) 발생"],
        ["블로그 목록", "/blog", "카테고리 필터·키워드 검색을 갖춘 콘텐츠 허브"],
        ["블로그 상세", "/blog/[slug]", "글별 SEO 메타·OG 태그, 조회수 집계"],
        ["서비스소개서", "/brochure·GNB 모달", "회사·담당자 정보 입력 시 소개서 PDF 메일 발송(리드 확보), 다운로드 클릭 추적, GA4 전환 이벤트(brochure_lead)"],
        ["법적 고지", "/privacy·/terms·/terms-applicant", "개인정보처리방침, 기업용·지원자용 이용약관. DB 기반 버전·시행일 관리"],
        ["공통 요소", "전 페이지", "섹션 인지형 GNB, FAQ 챗봇, 쿠키 동의 배너(GA4 Consent Mode v2 연동), 반응형 레이아웃, 404 페이지"],
    ],
    widths=[1.1, 1.5, 3.9],
)

P("수행 결과 이미지 (웹사이트 메인/서브 화면) :")
picture("landing-01-hero.png", "그림 1. 메인 랜딩 — 히어로 및 고객사 영역")
picture("landing-02-value.png", "그림 2. 메인 랜딩 — 핵심 가치 제안 섹션")
picture("landing-03-role.png", "그림 3. 메인 랜딩 — AI·사람 역할 분담(선별 퍼널) 섹션")
picture("landing-04-how.png", "그림 4. 메인 랜딩 — AI 리포트 소개 섹션")
picture("landing-05-proof.png", "그림 5. 메인 랜딩 — 검증된 효과(데이터 근거) 섹션")
picture("landing-06-voices.png", "그림 6. 메인 랜딩 — 고객 활용 사례 섹션")
picture("landing-08-final.png", "그림 7. 메인 랜딩 — 최종 상담 신청 CTA 및 푸터")
picture("apply.png", "그림 8. 도입 문의 페이지(/apply) — 상담 신청 폼")
picture("privacy.png", "그림 9. 개인정보처리방침 페이지(/privacy) — DB 기반 약관 렌더링")

H3("2) 관리자(CMS) 페이지")
P("콘텐츠·문의·약관·SEO를 코드 수정 없이 운영할 수 있는 자체 관리자 콘솔(/admin)을 구축함. "
  "Google 계정 로그인과 관리자 화이트리스트(admins) 이중 검증으로 접근을 제한함.")
table(
    ["메뉴", "기능"],
    [
        ["대시보드", "오늘·이번 주 도입문의, 공개 블로그 수 등 전체 현황 요약과 최근 활동 타임라인"],
        ["블로그", "블로그 글 등록·수정·삭제, 공개/비공개 전환, 에디터·커버 이미지 관리"],
        ["업데이트", "서비스 업데이트 소식 게시 관리"],
        ["FAQ", "웹사이트 FAQ·챗봇 문답 항목 관리"],
        ["소개서", "서비스 소개서 파일 관리(도입문의·웹사이트에서 제공되는 PDF)"],
        ["약관", "개인정보처리방침·이용약관 버전 관리 및 게시(시행일 관리)"],
        ["SEO", "페이지별 검색·공유 메타데이터(제목·설명·OG) 편집·적용"],
        ["도입문의", "접수된 도입문의 열람, 상담 상태 관리, UTM·유입 경로 등 마케팅 추적 정보 확인"],
        ["설정", "관리자 계정, GA4 측정 ID 등 사이트 기본 설정"],
    ],
    widths=[1.2, 5.3],
)
picture("admin-gate.png", "그림 10. 관리자 콘솔 로그인 게이트 — Google 계정 + 관리자 화이트리스트 검증")
picture("admin-dash.png", "그림 11. 관리자 — 대시보드 (도입문의·블로그·FAQ 현황 및 최근 활동, 고객 정보는 마스킹 처리)")
picture("admin-blog.png", "그림 12. 관리자 — 블로그 관리 (등록 41건·누적 조회 집계, 공개/비공개·수정·삭제)")
picture("admin-updates.png", "그림 13. 관리자 — 제품 업데이트 관리 (링크 공유형 비공개 페이지 운영)")
picture("admin-faq.png", "그림 14. 관리자 — FAQ 관리 (카테고리·노출 순서 관리)")
picture("admin-brochure.png", "그림 15. 관리자 — 서비스 소개서 관리 (PDF 교체 업로드, 신청 리드 16건 목록·CSV 다운로드, 고객 정보는 마스킹 처리)")
picture("admin-legal.png", "그림 16. 관리자 — 약관 관리 (개인정보처리방침·이용약관 버전/시행일 이력 관리)")
picture("admin-seo.png", "그림 17. 관리자 — SEO 메타데이터 관리 (페이지별 제목·설명·OG 초안 작성 및 적용)")
picture("admin-signups.png", "그림 18. 관리자 — 도입문의 관리 (상태·기간·규모 필터, 검색, CSV 다운로드, 고객 정보는 마스킹 처리)")
picture("admin-settings.png", "그림 19. 관리자 — 설정 (Google 계정 기반 관리자 추가, GA4 측정 ID 관리)")

# ══════════════════════════════════════════════════════════
# ② 마케팅 콘텐츠 기획 및 제작
# ══════════════════════════════════════════════════════════
H2("② 마케팅 콘텐츠 기획 및 제작")
P(f"과업 기간 동안 AI 면접·채용 검증·채용 트렌드·채용 자동화·HR 인사이트 5개 카테고리로 블로그 콘텐츠 "
  f"총 {len(posts)}건을 기획·제작하여 이 중 {len(pub)}건을 발행함(잔여 {len(drafts)}건은 발행 대기). "
  "전 콘텐츠에 자체 제작 커버 이미지와 검색엔진 최적화 메타데이터(제목·설명·OG)를 적용하고, "
  "웹사이트 블로그 채널(www.supercoder.co/blog)에 게재함.",
  bold_prefix="수행 요약 : ")

H3("1) 카테고리별 발행 현황")
cat_rows = [[c, f"{n}건"] for c, n in sorted(cats.items(), key=lambda x: -x[1])]
cat_rows.append(["합계", f"{len(pub)}건"])
table(["카테고리", "발행 콘텐츠"], cat_rows, widths=[2.5, 2.0])

H3("2) 발행 콘텐츠 목록")
rows = [[i + 1, p["created_at"][:10], p["category"], p["title"]] for i, p in enumerate(pub)]
table(["No", "발행일", "카테고리", "제목"], rows, widths=[0.4, 1.0, 1.0, 4.1], size=8)

P("수행 결과 이미지 (블로그 채널 및 게재 화면) :")
picture("blog-list.png", "그림 20. 블로그 목록 페이지(/blog) — 카테고리 필터·검색")
picture("blog-list-2.png", "그림 21. 블로그 목록 — 자체 제작 커버 이미지 기반 콘텐츠 카드")
picture("blog-post.png", "그림 22. 블로그 상세 페이지 — 게재 콘텐츠 예시")
picture("blog-post-2.png", "그림 23. 블로그 상세 — 본문 게재 화면")

# ══════════════════════════════════════════════════════════
# ③ 마케팅 채널 운영 및 성과 관리
# ══════════════════════════════════════════════════════════
H2("③ 마케팅 채널 운영 및 성과 관리")
P("Google Analytics(GA4)를 웹사이트에 연동하여 유입·참여·전환 지표를 측정하고, 도입문의(apply_lead)·"
  "소개서 신청(brochure_lead)을 GA4 주요 이벤트(전환)로 설정함. GA4 데이터 기반 Looker Studio 성과 대시보드"
  "(트래픽 현황·마케팅 전환 2페이지)를 구축하고, 매일 오전 8시 핵심 지표 요약이 자동 발송되는 일일 애널리틱스 "
  "리포트 체계를 운영함. 쿠키 동의 배너 도입 이후의 측정 공백을 보완하기 위해 쿠키리스 방식의 Vercel Analytics를 "
  "병행 도입함.",
  bold_prefix="수행 요약 : ")

H3("1) 웹사이트 트래픽 (GA4)")
table(
    ["기간", "세션 합계", "일평균 세션", "일 최대 세션", "비고"],
    [
        ["2026.07.08 ~ 07.30 (23일)", "796", "34.6", "66 (7/13)", "쿠키 동의 배너 도입 전 — 해외 봇 트래픽 일부 혼입"],
        ["2026.08.01 ~ 08.12 (12일)", "38", "3.2", "6 (8/8)", "쿠키 동의 기반 실측(동의 방문자만 집계)"],
    ],
    widths=[1.7, 0.9, 0.9, 0.9, 2.1],
)
P("※ 측정 기준 변경 안내 : 2026.07.29 개인정보 보호 강화를 위해 쿠키 동의 배너를, 08.07 Google Consent Mode v2를 "
  "도입함. 이후 GA4는 '동의한 방문자'만 집계하므로 8월 수치는 실제 방문의 하한선임. 별도 트래픽 품질 분석 결과 "
  "7월 수치에는 해외 클라우드발 봇 트래픽이 혼입되어 있었음(정상 세션 비율 약 46.5%). 이에 따라 국내·봇 제외 "
  "세그먼트 기반 모니터링 보고서를 구축하여 운영 중임.").runs[0].font.size = Pt(9)
picture("ga4-daily-sessions.png", "그림 24. GA4 일별 세션 추이 (2026.07.08~08.12, 일일 애널리틱스 리포트 집계) — 점선: 쿠키 동의 배너 도입 시점", width=6.4)

P("유입 채널은 Direct·Organic Search 중심이며, 기기별로는 데스크톱 비중(75~96%)이 우세함. "
  "도입문의·소개서 신청 리드는 접수 즉시 DB 저장과 담당자 메일 알림으로 전달되며, 관리자 콘솔에서 "
  "UTM·유입 경로와 함께 상담 상태를 관리함.")
picture("ga4-snapshot.png", "그림 25. GA4 보고서 홈 — 최근 30일 활성 사용자 413명·세션 630건 (캡처일 2026.08.13)")
picture("ga4-channels.png", "그림 26. GA4 트래픽 획득 보고서 — 세션 소스/매체별 추이 (2026.05.01~08.12)")
picture("looker-dashboard.png", "그림 27. Looker Studio 성과 대시보드 'Supercoder 웹 애널리틱스' — 트래픽 현황 페이지")

H3("2) 성과 대시보드 — 블로그 콘텐츠 조회수")
P(f"발행 콘텐츠 {len(pub)}건의 누적 조회수는 총 {total_views:,}회임(2026.08.13 기준, 웹사이트 자체 집계).")
by_views = sorted(pub, key=lambda p: -p["views"])
vrows = [[i + 1, p["title"], p["category"], p["created_at"][:10], f'{p["views"]}회'] for i, p in enumerate(by_views)]
vrows.append(["", "합계", "", "", f"{total_views:,}회"])
table(["순위", "제목", "카테고리", "발행일", "조회수"], vrows, widths=[0.4, 3.5, 1.0, 0.9, 0.7], size=8)

H3("3) 성과 관리 자동화 체계")
table(
    ["구성 요소", "내용"],
    [
        ["GA4 전환 이벤트", "도입문의(apply_lead)·소개서 신청(brochure_lead)을 주요 이벤트로 설정, 폼 제출 성공 시 자동 수집"],
        ["Looker Studio 대시보드", "GA4 연동 '트래픽 현황'·'마케팅 전환' 2페이지 구성 — 활성 사용자·세션·참여율, 채널·기기·페이지별 분석, 전환 이벤트 추이"],
        ["일일 애널리틱스 리포트", "매일 오전 8시(KST) 전일 핵심 지표(사용자·세션·참여율·전환·인기 페이지·유입 채널)와 자동 인사이트를 이메일 발송"],
        ["트래픽 품질 관리", "봇·내부 트래픽 분석 및 국내·봇 제외 세그먼트 모니터링 보고서 운영, 프리뷰 도메인 측정 차단"],
        ["보완 측정", "Vercel Analytics(쿠키리스) 병행 도입(2026.08), 리드 유입 경로 추적(UTM·클릭 ID·referrer)"],
    ],
    widths=[1.6, 4.9],
)

# ── 맺음말 ─────────────────────────────────────────────────
d.add_paragraph()
p = d.add_paragraph("위와 같이 과업 수행 결과를 보고합니다.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = d.add_paragraph("2026년 8월 31일")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = d.add_paragraph("홍보법인 동서남북 대표 유승훈 (인)")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.save(DOC)
print("saved:", DOC)
print("published:", len(pub), "drafts:", len(drafts), "total views:", total_views, "cats:", cats)
